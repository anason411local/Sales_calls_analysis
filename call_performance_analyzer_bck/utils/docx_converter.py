"""
Convert Markdown report to professional DOCX format
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re
from pathlib import Path
from typing import List


class ReportConverter:
    """Convert Markdown report to styled DOCX"""
    
    def __init__(self):
        self.doc = Document()
        self._setup_styles()
    
    def _setup_styles(self):
        """Setup custom styles for the document"""
        
        # Title style
        if 'Report Title' not in self.doc.styles:
            title_style = self.doc.styles.add_style('Report Title', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = 'Calibri'
            title_font.size = Pt(24)
            title_font.bold = True
            title_font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
            title_style.paragraph_format.space_after = Pt(12)
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Heading 1 style (already exists, just modify)
        heading1 = self.doc.styles['Heading 1']
        heading1.font.name = 'Calibri'
        heading1.font.size = Pt(18)
        heading1.font.color.rgb = RGBColor(0, 51, 102)
        
        # Heading 2 style
        heading2 = self.doc.styles['Heading 2']
        heading2.font.name = 'Calibri'
        heading2.font.size = Pt(14)
        heading2.font.color.rgb = RGBColor(0, 102, 204)
        
        # Quote style for verbatim examples
        if 'Verbatim Quote' not in self.doc.styles:
            quote_style = self.doc.styles.add_style('Verbatim Quote', WD_STYLE_TYPE.PARAGRAPH)
            quote_font = quote_style.font
            quote_font.name = 'Courier New'
            quote_font.size = Pt(10)
            quote_font.italic = True
            quote_style.paragraph_format.left_indent = Inches(0.5)
            quote_style.paragraph_format.space_before = Pt(6)
            quote_style.paragraph_format.space_after = Pt(6)
    
    def convert_markdown_to_docx(self, markdown_file: str, output_file: str):
        """
        Convert markdown report to DOCX
        
        Args:
            markdown_file: Path to markdown file
            output_file: Path to output DOCX file
        """
        # Read markdown content
        with open(markdown_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Parse and convert
        self._parse_content(content)
        
        # Save document
        self.doc.save(output_file)
        print(f"[SUCCESS] Report converted to DOCX: {output_file}")
    
    def _parse_content(self, content: str):
        """Parse markdown content and add to document"""
        
        lines = content.split('\n')
        i = 0
        in_table = False
        table_data = []
        in_code_block = False
        code_lines = []
        
        while i < len(lines):
            line = lines[i]
            
            # Skip empty lines (but add spacing)
            if not line.strip():
                if not in_table and not in_code_block:
                    # Add small spacing
                    pass
                i += 1
                continue
            
            # Code blocks (verbatim quotes)
            if line.strip().startswith('```'):
                if in_code_block:
                    # End of code block
                    self._add_code_block(code_lines)
                    code_lines = []
                    in_code_block = False
                else:
                    # Start of code block
                    in_code_block = True
                i += 1
                continue
            
            if in_code_block:
                code_lines.append(line)
                i += 1
                continue
            
            # Title (##)
            if line.startswith('## '):
                title = line[3:].strip()
                p = self.doc.add_paragraph(title, style='Report Title')
                i += 1
                continue
            
            # Heading 1 (###)
            if line.startswith('### '):
                heading = line[4:].strip()
                self.doc.add_heading(heading, level=1)
                i += 1
                continue
            
            # Heading 2 (####)
            if line.startswith('#### '):
                heading = line[5:].strip()
                self.doc.add_heading(heading, level=2)
                i += 1
                continue
            
            # Tables
            if '|' in line and not in_table:
                # Start of table
                in_table = True
                table_data = []
                table_data.append(line)
                i += 1
                continue
            
            if in_table:
                if '|' in line:
                    # Skip separator lines
                    if not line.strip().replace('|', '').replace('-', '').replace(':', '').replace(' ', ''):
                        i += 1
                        continue
                    table_data.append(line)
                    i += 1
                    continue
                else:
                    # End of table
                    self._add_table(table_data)
                    table_data = []
                    in_table = False
                    continue
            
            # Bold text with asterisks
            if '**' in line:
                self._add_formatted_paragraph(line)
                i += 1
                continue
            
            # Bullet points
            if line.strip().startswith('*   ') or line.strip().startswith('- '):
                bullet_text = line.strip()[4:] if line.strip().startswith('*   ') else line.strip()[2:]
                self._add_bullet(bullet_text)
                i += 1
                continue
            
            # Numbered lists
            if re.match(r'^\d+\.\s+', line.strip()):
                list_text = re.sub(r'^\d+\.\s+', '', line.strip())
                self._add_numbered_item(list_text)
                i += 1
                continue
            
            # Horizontal rule
            if line.strip() == '---':
                self.doc.add_paragraph('_' * 80)
                i += 1
                continue
            
            # Regular paragraph
            if line.strip():
                self._add_formatted_paragraph(line.strip())
            
            i += 1
        
        # Handle any remaining table
        if in_table and table_data:
            self._add_table(table_data)
    
    def _add_formatted_paragraph(self, text: str):
        """Add paragraph with bold formatting"""
        p = self.doc.add_paragraph()
        
        # Split by bold markers
        parts = re.split(r'(\*\*.*?\*\*)', text)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Bold text
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                # Regular text
                p.add_run(part)
    
    def _add_bullet(self, text: str):
        """Add bullet point"""
        p = self.doc.add_paragraph(style='List Bullet')
        
        # Handle bold text in bullets
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)
    
    def _add_numbered_item(self, text: str):
        """Add numbered list item"""
        p = self.doc.add_paragraph(style='List Number')
        
        # Handle bold text
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)
    
    def _add_code_block(self, lines: List[str]):
        """Add code block (verbatim quote)"""
        if not lines:
            return
        
        for line in lines:
            p = self.doc.add_paragraph(line, style='Verbatim Quote')
    
    def _add_table(self, table_data: List[str]):
        """Add table to document"""
        if not table_data:
            return
        
        # Parse table data
        rows = []
        for line in table_data:
            cells = [cell.strip() for cell in line.split('|')[1:-1]]  # Remove empty first/last
            if cells:
                rows.append(cells)
        
        if not rows:
            return
        
        # Create table
        table = self.doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = 'Light Grid Accent 1'
        
        # Fill table
        for i, row_data in enumerate(rows):
            row = table.rows[i]
            for j, cell_data in enumerate(row_data):
                cell = row.cells[j]
                cell.text = cell_data
                
                # Bold header row
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
        
        # Add spacing after table
        self.doc.add_paragraph()


def convert_report_to_docx(markdown_path: str, docx_path: str = None):
    """
    Convert markdown report to DOCX
    
    Args:
        markdown_path: Path to markdown report
        docx_path: Output path for DOCX (optional, auto-generated if not provided)
    """
    markdown_path = Path(markdown_path)
    
    if docx_path is None:
        docx_path = markdown_path.parent / f"{markdown_path.stem}.docx"
    
    converter = ReportConverter()
    converter.convert_markdown_to_docx(str(markdown_path), str(docx_path))
    
    return str(docx_path)


if __name__ == "__main__":
    # Example usage
    import sys
    
    if len(sys.argv) > 1:
        markdown_file = sys.argv[1]
        docx_file = sys.argv[2] if len(sys.argv) > 2 else None
        convert_report_to_docx(markdown_file, docx_file)
    else:
        print("Usage: python docx_converter.py <markdown_file> [output_docx]")

